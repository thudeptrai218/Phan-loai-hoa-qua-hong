from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def clear_document(document):
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def setup_document(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)


def add_center(document, text, size=14, bold=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
    return paragraph


def add_para(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Inches(0.3)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_bullet(document, text):
    add_para(document, "- " + text)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def add_code_excerpt(document, title, path, max_lines=34):
    add_heading(document, title, 2)
    file_path = Path(path)
    if not file_path.exists():
        add_para(document, f"Khong tim thay file {path}.")
        return
    lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()[:max_lines]
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_common_paragraphs(document, paragraphs):
    for paragraph in paragraphs:
        add_para(document, paragraph)


def main():
    if len(sys.argv) < 3:
        raise SystemExit("Usage: python generate_three_chapter_report.py TEMPLATE_DOCX OUTPUT_DOCX")

    template_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    if not template_path.exists():
        raise FileNotFoundError(template_path)

    document = Document(str(template_path))
    clear_document(document)
    setup_document(document)

    add_center(document, "BAO CAO DO AN", 18)
    add_center(document, "HE THONG PHAN LOAI QUA TUOI VA QUA HONG", 18)
    add_center(document, "SU DUNG YOLO - DROIDCAM - FLASK - ESP8266 - ARDUINO UNO", 13)
    document.add_paragraph()
    add_center(document, "Sinh vien thuc hien: Doan Tuan Nam", 13)
    add_center(document, "Ngay tao bao cao: " + datetime.now().strftime("%d/%m/%Y"), 12, False)
    document.add_page_break()

    add_heading(document, "MUC LUC", 1)
    add_para(document, "CHUONG 1. TONG QUAN DE TAI VA CO SO LY THUYET")
    add_para(document, "CHUONG 2. PHAN TICH, THIET KE VA XAY DUNG HE THONG")
    add_para(document, "CHUONG 3. THUC NGHIEM, KIEM THU VA DANH GIA")
    add_para(document, "KET LUAN")
    add_para(document, "TAI LIEU THAM KHAO")
    add_para(document, "PHU LUC")
    document.add_page_break()

    add_heading(document, "CHUONG 1. TONG QUAN DE TAI VA CO SO LY THUYET", 1)
    sections = [
        ("1.1 Ly do chon de tai", [
            "Trong qua trinh bao quan va phan phoi nong san, viec phan loai trai cay theo tinh trang tuoi va hong co vai tro quan trong. Neu qua hong khong duoc tach ra kip thoi, chat luong lo hang co the bi anh huong va gay that thoat kinh te.",
            "Phan loai bang thu cong ton nhieu thoi gian, phu thuoc vao kinh nghiem cua nguoi lao dong va kho duy tri do on dinh khi so luong trai cay lon. Vi vay, viec ung dung thi giac may tinh va dieu khien tu dong la huong tiep can phu hop.",
            "De tai nay ket hop mo hinh YOLO de nhan dien qua tuoi va qua hong, ung dung web Flask de hien thi ket qua, DroidCam lam nguon camera, ESP8266 lam cau noi Wi-Fi va Arduino Uno dieu khien servo gat san pham."
        ]),
        ("1.2 Muc tieu nghien cuu", [
            "Muc tieu cua de tai la xay dung mot he thong mau co kha nang nhan dien hai trang thai Qua Hong va Qua Tuoi trong thoi gian thuc. Ket qua nhan dien duoc dung de hien thi tren web va dieu khien servo phan loai.",
            "He thong can dam bao cac thanh phan lam viec phoi hop: camera truyen hinh anh, model YOLO phan tich frame, Flask dieu phoi web va lenh phan cung, ESP8266 nhan lenh HTTP, Arduino Uno dieu khien hai servo.",
        ]),
        ("1.3 Doi tuong va pham vi nghien cuu", [
            "Doi tuong nghien cuu la bai toan nhan dien trai cay theo tinh trang hong hoac tuoi trong anh camera. Pham vi thuc hien tap trung vao mo hinh demo, chua huong den day chuyen cong nghiep hoan chinh.",
            "He thong su dung tap du lieu tu Roboflow voi hai nhan Qua Hong va Qua Tuoi. Camera lay tu DroidCam, server chay tren laptop va phan cung gom ESP8266, Arduino Uno cung hai servo."
        ]),
        ("1.4 Phuong phap thuc hien", [
            "De tai duoc thuc hien theo quy trinh: thu thap du lieu, gan nhan tren Roboflow, export dataset YOLOv8, train model bang Ultralytics, tich hop model vao Flask, xay dung giao dien web, lap mach ESP8266 va Arduino Uno, cuoi cung kiem thu toan he thong.",
            "Qua trinh phat trien duoc chia thanh tung module de de kiem tra rieng. Khi tung module hoat dong on dinh moi ghep thanh he thong hoan chinh."
        ]),
        ("1.5 Co so ly thuyet ve YOLO", [
            "YOLO la nhom mo hinh phat hien doi tuong thoi gian thuc. Mo hinh du doan dong thoi hop gioi han, class va do tin cay cua doi tuong trong anh. Uu diem cua YOLO la toc do nhanh, phu hop voi ung dung camera truc tiep.",
            "Trong project, model YOLO duoc train lai tren hai lop Qua Hong va Qua Tuoi. Khi server nhan frame tu camera, model tra ve danh sach box. Cac box dat nguong tin cay se duoc ve len anh va duoc dung de ra quyet dinh dieu khien servo."
        ]),
        ("1.6 Co so ly thuyet ve Flask va xu ly video", [
            "Flask la framework Python nhe, phu hop de tao nhanh web server cho cac ung dung demo AI. Trong he thong nay, Flask phuc vu giao dien web, route video_feed, route history va route hardware_status.",
            "Video duoc gui ve trinh duyet theo dang multipart JPEG. Moi frame duoc OpenCV doc tu DroidCam, dua qua YOLO, ve ket qua roi ma hoa thanh JPEG."
        ]),
        ("1.7 Co so ly thuyet ve ESP8266 va Arduino Uno", [
            "ESP8266 la module vi dieu khien co Wi-Fi, co the tao HTTP server de nhan lenh tu Flask. Sau khi nhan lenh servo=1 hoac servo=2, ESP8266 gui ky tu tuong ung qua Serial cho Arduino Uno.",
            "Arduino Uno nhan lenh Serial va dieu khien servo bang thu vien Servo.h. Hai servo duoc gan voi hai tinh trang: servo 1 cho qua hong, servo 2 cho qua tuoi."
        ]),
    ]
    for title, paragraphs in sections:
        add_heading(document, title, 2)
        add_common_paragraphs(document, paragraphs)

    add_heading(document, "1.8 Cong nghe su dung", 2)
    add_table(document, ["Thanh phan", "Vai tro"], [
        ["Python", "Xu ly AI, Flask server va script train"],
        ["Ultralytics YOLO", "Nhan dien Qua Hong va Qua Tuoi"],
        ["Roboflow", "Quan ly va export dataset"],
        ["OpenCV", "Doc camera va ve ket qua"],
        ["Flask", "Xay dung web dashboard"],
        ["DroidCam", "Nguon camera tu dien thoai"],
        ["ESP8266", "Nhan lenh HTTP tu server"],
        ["Arduino Uno", "Dieu khien hai servo"],
    ])
    document.add_page_break()

    add_heading(document, "CHUONG 2. PHAN TICH, THIET KE VA XAY DUNG HE THONG", 1)
    sections = [
        ("2.1 Yeu cau bai toan", [
            "He thong can nhan dien dung hai loai doi tuong: Qua Hong va Qua Tuoi. Cac doi tuong khac khong nam trong hai class cua model se khong duoc xu ly. Ket qua nhan dien phai duoc hien thi tren web va dong thoi sinh lenh dieu khien phan cung.",
            "Ngoai nhan dien, he thong can co kha nang thong bao trang thai ket noi ESP8266, tranh spam lenh servo va khong lam treo luong video khi phan cung mat ket noi."
        ]),
        ("2.2 Kien truc tong the", [
            "Kien truc he thong gom bon lop: lop camera, lop AI, lop web va lop phan cung. DroidCam cung cap video cho laptop. Flask lay frame, goi YOLO de nhan dien va day lenh vao hang doi servo. Worker nen gui lenh HTTP den ESP8266. ESP8266 gui Serial cho Arduino Uno de gat servo.",
            "Thiet ke nay giup tach luong video khoi luong phan cung. Neu ESP8266 cham hoac mat ket noi, DroidCam khong bi dung do request HTTP da duoc dua sang worker rieng."
        ]),
        ("2.3 Thiet ke dataset va train model", [
            "Dataset duoc quan ly tren Roboflow voi hai class: 0 la Qua Hong va 1 la Qua Tuoi. Thu tu class nay duoc dong bo voi server.py de anh xa class 0 thanh servo=1 va class 1 thanh servo=2.",
            "File train_model.py chon data.yaml phu hop va train model YOLO. Model sau train duoc dat tai runs/detect/rotten_fruit/weights/best.pt. Server chi load file model nay de tranh nhan dien sai tu model mac dinh."
        ]),
        ("2.4 Thiet ke server Flask", [
            "Server Flask khoi tao YOLO, doc camera, ve ket qua va phuc vu web. Cac route chinh gom /, /video_feed, /history va /hardware_status. Route /hardware_status cho phep giao dien biet lenh gan nhat, phan hoi OK hoac loi ket noi ESP8266.",
            "Cac thong so quan trong nhu DROIDCAM_IP, ESP8266_URL, SERVO_COOLDOWN_SECONDS va ENABLE_HARDWARE co the cau hinh qua bien moi truong hoac script run_all.ps1."
        ]),
        ("2.5 Thiet ke giao dien web", [
            "Giao dien web duoc chia thanh khu vuc video va khu vuc thong tin. Khu vuc video hien thi frame nhan dien truc tiep. Khu vuc thong tin hien thi thong ke so qua tuoi, qua hong, lich su quet va trang thai phan cung.",
            "Phan trang thai phan cung giup nguoi dung biet ESP8266 da phan hoi hay chua. Khi co loi timeout, loi duoc hien thi tren web de phuc vu debug."
        ]),
        ("2.6 Thiet ke giao tiep ESP8266", [
            "ESP8266 co endpoint /ping de kiem tra ket noi va endpoint /control de nhan lenh servo. Khi nhan /control?servo=1, ESP8266 in so 1 qua Serial. Khi nhan /control?servo=2, ESP8266 in so 2.",
            "Firmware ESP8266 duoc cau hinh de ket noi Wi-Fi chinh. Neu khong ket noi duoc, ESP8266 co che do Wi-Fi du phong PHAN_LOAI_TRAI_CAY de test doc lap."
        ]),
        ("2.7 Thiet ke Arduino Uno va servo", [
            "Arduino Uno nhan ky tu Serial tu ESP8266. Lenh 1 kich hoat servoHong o chan D9, lenh 2 kich hoat servoTuoi o chan D10. Sau khi gat, servo quay ve goc ban dau.",
            "Servo tuoi co do tre SERVO_TUOI_DELAY de phu hop voi vi tri vat tren bang chuyen. Cac gia tri goc va thoi gian co the dieu chinh trong code."
        ]),
        ("2.8 Co che chong spam va tang do on dinh", [
            "Chuong trinh chi gui lenh khi class nhan dien on dinh trong mot so frame lien tiep. Moi servo co cooldown rieng de tranh gat lien tuc. Neu ESP8266 loi, chuong trinh tam dung gui lai trong thoi gian cau hinh.",
            "Lenh servo duoc dua vao queue nen. Queue chi giu lenh moi nhat, tranh tinh trang ton dong lenh cu lam servo hoat dong sai thoi diem."
        ]),
    ]
    for title, paragraphs in sections:
        add_heading(document, title, 2)
        add_common_paragraphs(document, paragraphs)

    add_heading(document, "2.9 Bang anh xa class va servo", 2)
    add_table(document, ["Class", "Nhan", "Trang thai", "Lenh ESP8266", "Chan Arduino"], [
        ["0", "Qua Hong", "Hong", "servo=1", "D9"],
        ["1", "Qua Tuoi", "Tuoi", "servo=2", "D10"],
    ])
    document.add_page_break()

    add_heading(document, "CHUONG 3. THUC NGHIEM, KIEM THU VA DANH GIA", 1)
    sections = [
        ("3.1 Moi truong thuc nghiem", [
            "He thong duoc chay tren laptop Windows, code Python su dung Flask, OpenCV, Ultralytics va requests. Phan cung gom ESP8266 NodeMCU, Arduino Uno va hai servo. Camera dung DroidCam voi IP mac dinh da cau hinh trong project.",
            "Qua trinh upload phan cung thuc hien bang PlatformIO trong VS Code. Project co script upload_esp8266.ps1, upload_uno.ps1 va run_all.ps1 de giam thao tac thu cong."
        ]),
        ("3.2 Kiem thu model nhan dien", [
            "Model duoc train tu dataset Roboflow va luu thanh best.pt. Khi load model, he thong xac nhan hai nhan la Qua Hong va Qua Tuoi. Khi dua anh mau vao, model tra ve box va confidence cho class tuong ung.",
            "Do dataset hien tai con nho, ket qua nhan dien trong dieu kien thuc te can duoc cai thien bang cach bo sung anh moi va train lai."
        ]),
        ("3.3 Kiem thu DroidCam", [
            "DroidCam duoc kiem thu bang viec mo URL video tren server. Neu camera khong ket noi duoc, server hien frame thong bao loi thay vi lam web dung dot ngot. Khi ESP8266 cham, viec gui lenh da tach sang worker nen nen luong camera khong bi dung.",
            "Neu laptop ket noi vao Wi-Fi du phong cua ESP8266, DroidCam qua router co the bi mat ket noi. Do do trong van hanh on dinh nen de DroidCam va laptop cung mang Wi-Fi, hoac dung DroidCam qua USB."
        ]),
        ("3.4 Kiem thu ESP8266", [
            "ESP8266 duoc kiem thu qua endpoint /ping. Neu trinh duyet hoac PowerShell truy cap duoc /ping va nhan duoc chuoi ESP8266 OK, nghia la ket noi mang da thanh cong. Sau do test /control?servo=1 va /control?servo=2 de kiem tra lenh.",
            "Trong qua trinh phat trien da gap loi timeout do sai IP va loi upload do board chua vao bootloader. Project da bo sung script va huong dan de khac phuc."
        ]),
        ("3.5 Kiem thu Arduino Uno va servo", [
            "Arduino duoc kiem thu rieng bang cach gui ky tu 1 hoac 2 qua Serial Monitor. Neu servo D9 va D10 quay dung, phan Arduino hoat dong. Sau do moi noi ESP8266 TX sang Arduino RX de kiem thu chuoi ESP8266 -> Arduino.",
            "Nguon servo can du dong va GND phai noi chung. Neu servo khong quay, can kiem tra nguon, day signal va trang thai code Arduino."
        ]),
        ("3.6 Kiem thu toan he thong", [
            "Khi toan he thong chay, nguoi dung mo web Flask, dua trai cay vao camera va quan sat ket qua. Neu phat hien Qua Hong, giao dien hien khung do va server gui servo=1. Neu phat hien Qua Tuoi, giao dien hien khung xanh va server gui servo=2.",
            "Giao dien phan cung cap nhat lenh gan nhat, phan hoi OK va loi neu ESP8266 khong phan hoi. Day la cong cu quan trong de xac dinh he thong dang loi o camera, model, server hay phan cung."
        ]),
        ("3.7 Ket qua dat duoc", [
            "Project da hoan thanh cac chuc nang chinh: train model Roboflow, nhan dien hai class, hien thi web, dieu khien ESP8266 va Arduino Uno, chong spam servo va bo sung script chay mot lenh.",
            "He thong co the dung lam mo hinh demo AIoT cho bai toan phan loai trai cay theo tinh trang tuoi va hong."
        ]),
        ("3.8 Han che va huong phat trien", [
            "Han che hien tai la dataset con it anh, chua co day du valid/test doc lap va mo hinh co khi bang chuyen chua duoc toi uu. Ket noi Wi-Fi giua DroidCam, laptop va ESP8266 cung can duoc bo tri on dinh.",
            "Huong phat trien la mo rong dataset, bo sung cam bien phat hien vat, thay camera bang camera USB hoac camera cong nghiep, cai tien co khi gat va luu lich su vao co so du lieu."
        ]),
    ]
    for title, paragraphs in sections:
        add_heading(document, title, 2)
        add_common_paragraphs(document, paragraphs)

    add_heading(document, "3.9 Bang tong hop loi va cach xu ly", 2)
    add_table(document, ["Loi", "Nguyen nhan", "Cach xu ly"], [
        ["ESP8266 timeout", "Sai IP hoac khac mang", "Dung /ping, kiem tra IP, dung auto-discover"],
        ["DroidCam dung", "Request phan cung chan luong video", "Dung queue va worker nen"],
        ["Upload ESP8266 fail", "Board chua vao bootloader hoac COM bi giu", "Dong Serial Monitor, thao RX/TX, giu FLASH/BOOT"],
        ["Servo khong quay", "Nguon yeu hoac sai day", "Dung nguon 5V ngoai va noi chung GND"],
    ])
    document.add_page_break()

    add_heading(document, "KET LUAN", 1)
    add_para(document, "Bao cao da trinh bay qua trinh xay dung he thong phan loai qua tuoi va qua hong su dung YOLO, DroidCam, Flask, ESP8266 va Arduino Uno. He thong the hien duoc quy trinh hoan chinh tu du lieu den hanh dong vat ly cua servo.")
    add_para(document, "Ket qua dat duoc cho thay mo hinh AI co the tich hop voi phan cung dieu khien trong mot ung dung AIoT. De tai co kha nang mo rong thanh he thong phan loai tu dong quy mo lon hon neu duoc bo sung du lieu va cai tien co khi.")
    document.add_page_break()

    add_heading(document, "TAI LIEU THAM KHAO", 1)
    for item in [
        "Ultralytics YOLO Documentation.",
        "Flask Documentation.",
        "OpenCV Documentation.",
        "Roboflow Documentation.",
        "ESP8266 Arduino Core Documentation.",
        "Arduino Servo Library Documentation.",
    ]:
        add_bullet(document, item)
    document.add_page_break()

    add_heading(document, "PHU LUC", 1)
    add_code_excerpt(document, "Phu luc A. Trich doan server.py", "server.py")
    add_code_excerpt(document, "Phu luc B. Trich doan ESP8266", "src/esp8266_node/main.cpp")
    add_code_excerpt(document, "Phu luc C. Trich doan Arduino Uno", "src/arduino_uno/main.cpp")
    add_code_excerpt(document, "Phu luc D. Trich doan run_all.ps1", "run_all.ps1")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
