from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def clear_document(document):
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def setup_document(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)


def center(document, text, size=14, bold=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"


def para(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Inches(0.3)
    paragraph.paragraph_format.line_spacing = 1.25


def bullet(document, text):
    para(document, "- " + text)


def table(document, headers, rows):
    tbl = document.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for index, header in enumerate(headers):
        tbl.rows[0].cells[index].text = header
    for row in rows:
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def code_excerpt(document, title, path, max_lines=32):
    heading(document, title, 2)
    file_path = Path(path)
    if not file_path.exists():
        para(document, f"Không tìm thấy tệp {path}.")
        return
    lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()[:max_lines]
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_section(document, title, paragraphs):
    heading(document, title, 2)
    for text in paragraphs:
        para(document, text)


def main():
    if len(sys.argv) < 3:
        raise SystemExit("Usage: python generate_three_chapter_report_vi.py TEMPLATE_DOCX OUTPUT_DOCX")

    template_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    if not template_path.exists():
        raise FileNotFoundError(template_path)

    document = Document(str(template_path))
    clear_document(document)
    setup_document(document)

    center(document, "BÁO CÁO ĐỒ ÁN", 18)
    center(document, "HỆ THỐNG PHÂN LOẠI QUẢ TƯƠI VÀ QUẢ HỎNG", 18)
    center(document, "SỬ DỤNG YOLO - DROIDCAM - FLASK - ESP8266 - ARDUINO UNO", 13)
    document.add_paragraph()
    center(document, "Sinh viên thực hiện: Đoàn Tuấn Nam", 13)
    center(document, "Ngày tạo báo cáo: " + datetime.now().strftime("%d/%m/%Y"), 12, False)
    document.add_page_break()

    heading(document, "MỤC LỤC", 1)
    para(document, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI VÀ CƠ SỞ LÝ THUYẾT")
    para(document, "CHƯƠNG 2. PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG")
    para(document, "CHƯƠNG 3. THỰC NGHIỆM, KIỂM THỬ VÀ ĐÁNH GIÁ")
    para(document, "KẾT LUẬN")
    para(document, "TÀI LIỆU THAM KHẢO")
    para(document, "PHỤ LỤC")
    document.add_page_break()

    heading(document, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI VÀ CƠ SỞ LÝ THUYẾT", 1)
    chapter_1 = [
        ("1.1 Lý do chọn đề tài", [
            "Trong quá trình bảo quản và phân phối nông sản, việc phân loại trái cây theo tình trạng tươi hoặc hỏng có ý nghĩa quan trọng. Nếu quả hỏng không được phát hiện và tách ra kịp thời, chất lượng của cả lô hàng có thể bị ảnh hưởng, gây thất thoát kinh tế và làm giảm độ tin cậy của quy trình phân phối.",
            "Phân loại thủ công phụ thuộc nhiều vào kinh nghiệm của người lao động, tốc độ xử lý không ổn định và dễ xảy ra sai sót khi số lượng sản phẩm lớn. Vì vậy, việc ứng dụng thị giác máy tính kết hợp với hệ thống điều khiển tự động là hướng tiếp cận phù hợp với xu thế nông nghiệp thông minh.",
            "Đề tài này xây dựng một mô hình tích hợp giữa phần mềm và phần cứng: mô hình YOLO dùng để nhận diện quả tươi và quả hỏng, DroidCam cung cấp hình ảnh, Flask hiển thị kết quả trên giao diện web, ESP8266 nhận lệnh qua Wi-Fi và Arduino Uno điều khiển servo để thực hiện thao tác phân loại."
        ]),
        ("1.2 Mục tiêu nghiên cứu", [
            "Mục tiêu chính của đề tài là xây dựng một hệ thống có khả năng nhận diện hai trạng thái của trái cây trong thời gian thực, gồm quả tươi và quả hỏng. Kết quả nhận diện được hiển thị trực tiếp trên giao diện web và được sử dụng để điều khiển cơ cấu servo.",
            "Hệ thống cần đảm bảo các chức năng cơ bản: thu nhận hình ảnh từ camera, xử lý bằng mô hình YOLO, hiển thị video đã vẽ khung nhận diện, lưu lịch sử quét, gửi lệnh đến ESP8266 và điều khiển Arduino Uno gạt trái cây về đúng vị trí.",
            "Ngoài chức năng nhận diện, đề tài cũng hướng đến việc xây dựng một quy trình triển khai dễ sử dụng, có thể chạy bằng script, có khả năng kiểm tra trạng thái phần cứng và có cơ chế chống gửi lệnh liên tục khi mô hình nhận diện nhiều khung hình liên tiếp."
        ]),
        ("1.3 Đối tượng và phạm vi nghiên cứu", [
            "Đối tượng nghiên cứu của đề tài là bài toán phát hiện và phân loại trái cây theo trạng thái tươi hoặc hỏng dựa trên hình ảnh camera. Trong phạm vi đồ án, hệ thống được xây dựng ở mức mô hình thử nghiệm, phục vụ trình diễn nguyên lý hoạt động và khả năng tích hợp AIoT.",
            "Phạm vi phần mềm bao gồm huấn luyện mô hình YOLO bằng dữ liệu Roboflow, xây dựng server Flask, xử lý video DroidCam và giao diện web. Phạm vi phần cứng bao gồm ESP8266, Arduino Uno và hai servo dùng để mô phỏng thao tác gạt sản phẩm.",
            "Đề tài chưa đi sâu vào thiết kế cơ khí công nghiệp, tốc độ băng chuyền thực tế hoặc hệ thống cấp liệu tự động. Các phần này có thể được phát triển thêm trong các giai đoạn tiếp theo."
        ]),
        ("1.4 Phương pháp thực hiện", [
            "Quy trình thực hiện đề tài bắt đầu từ việc chuẩn bị dữ liệu ảnh, gắn nhãn hai lớp quả tươi và quả hỏng trên Roboflow, sau đó xuất dữ liệu theo định dạng YOLOv8. Dữ liệu được dùng để huấn luyện mô hình bằng thư viện Ultralytics.",
            "Sau khi có mô hình best.pt, chương trình Flask được xây dựng để đọc video từ DroidCam, đưa từng khung hình qua mô hình YOLO, vẽ kết quả và hiển thị lên giao diện web. Khi phát hiện ổn định một trạng thái, server gửi lệnh điều khiển tới ESP8266.",
            "Phần cứng được kiểm thử theo từng lớp: kiểm tra ESP8266 bằng endpoint /ping, kiểm tra Arduino Uno bằng Serial, kiểm tra servo bằng lệnh 1 và 2, sau đó mới kết nối thành hệ thống hoàn chỉnh."
        ]),
        ("1.5 Cơ sở lý thuyết về YOLO", [
            "YOLO là nhóm mô hình phát hiện đối tượng thời gian thực. Mô hình có khả năng dự đoán đồng thời vị trí đối tượng, lớp đối tượng và độ tin cậy trong mỗi khung hình. Đặc điểm này phù hợp với các bài toán cần xử lý camera trực tiếp.",
            "Trong hệ thống, YOLO nhận đầu vào là ảnh từ DroidCam và trả về các bounding box tương ứng với hai lớp Qua Hong và Qua Tuoi. Mỗi kết quả gồm class id, tọa độ khung và confidence. Server chỉ xử lý các kết quả đạt ngưỡng tin cậy đã cấu hình.",
            "Việc sử dụng YOLO giúp hệ thống không chỉ biết trong ảnh có quả hay không, mà còn xác định vùng xuất hiện của quả để hiển thị trực quan trên giao diện."
        ]),
        ("1.6 Cơ sở lý thuyết về Flask và xử lý video", [
            "Flask là một framework Python nhẹ, phù hợp cho các ứng dụng web thử nghiệm và hệ thống demo AI. Trong đề tài, Flask đóng vai trò trung tâm điều phối giữa mô hình nhận diện, giao diện người dùng và phần cứng.",
            "Luồng video được truyền tới trình duyệt theo dạng multipart JPEG. Mỗi khung hình được đọc bằng OpenCV, xử lý bởi YOLO, vẽ nhãn và mã hóa lại thành ảnh JPEG trước khi gửi về web.",
            "Các route chính gồm trang chủ, video_feed, history và hardware_status. Cách tách route này giúp giao diện có thể cập nhật lịch sử và trạng thái phần cứng mà không cần tải lại toàn bộ trang."
        ]),
        ("1.7 Cơ sở lý thuyết về ESP8266 và Arduino Uno", [
            "ESP8266 là vi điều khiển có Wi-Fi tích hợp, có thể đóng vai trò web server nhỏ để nhận lệnh từ máy tính. Trong đề tài, ESP8266 nhận lệnh HTTP từ Flask và chuyển thành ký tự Serial gửi cho Arduino Uno.",
            "Arduino Uno chịu trách nhiệm điều khiển hai servo. Khi nhận ký tự 1, Arduino kích hoạt servo gạt quả hỏng. Khi nhận ký tự 2, Arduino kích hoạt servo gạt quả tươi. Cách phân chia nhiệm vụ này giúp ESP8266 tập trung vào giao tiếp mạng, còn Arduino tập trung vào điều khiển cơ cấu chấp hành.",
            "Việc kết nối ESP8266 với Arduino cần chú ý mức điện áp logic. Arduino Uno dùng mức 5V, trong khi ESP8266 dùng mức 3.3V, do đó đường tín hiệu từ Arduino TX sang ESP8266 RX cần có mạch chia áp nếu sử dụng."
        ]),
    ]
    for title, paragraphs in chapter_1:
        add_section(document, title, paragraphs)
    heading(document, "1.8 Công nghệ sử dụng", 2)
    table(document, ["Thành phần", "Vai trò trong hệ thống"], [
        ["Python", "Xử lý AI, chạy Flask server và các script hỗ trợ"],
        ["Ultralytics YOLO", "Nhận diện quả tươi và quả hỏng"],
        ["Roboflow", "Quản lý dữ liệu, gắn nhãn và xuất dataset"],
        ["OpenCV", "Đọc camera, xử lý khung hình và vẽ kết quả"],
        ["Flask", "Xây dựng giao diện web và API trạng thái"],
        ["DroidCam", "Cung cấp luồng camera từ điện thoại"],
        ["ESP8266", "Nhận lệnh HTTP và gửi Serial cho Arduino"],
        ["Arduino Uno", "Điều khiển hai servo phân loại"],
    ])
    document.add_page_break()

    heading(document, "CHƯƠNG 2. PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG", 1)
    chapter_2 = [
        ("2.1 Yêu cầu bài toán", [
            "Hệ thống cần nhận diện đúng hai lớp đối tượng: Qua Hong và Qua Tuoi. Các đối tượng không thuộc hai lớp này sẽ không được xử lý để tránh điều khiển nhầm phần cứng.",
            "Kết quả nhận diện cần được hiển thị trên giao diện web theo thời gian thực. Ngoài ra, hệ thống phải gửi lệnh điều khiển tương ứng tới ESP8266 và Arduino Uno để mô phỏng thao tác phân loại.",
            "Một yêu cầu quan trọng khác là chương trình không được bị treo khi ESP8266 phản hồi chậm hoặc mất kết nối. Vì vậy, việc gửi lệnh phần cứng được tách sang một worker nền thay vì thực hiện trực tiếp trong vòng đọc camera."
        ]),
        ("2.2 Kiến trúc tổng thể", [
            "Kiến trúc hệ thống gồm bốn lớp chính: lớp thu nhận hình ảnh, lớp xử lý AI, lớp giao diện web và lớp điều khiển phần cứng. DroidCam thuộc lớp thu nhận hình ảnh, YOLO thuộc lớp xử lý AI, Flask thuộc lớp giao diện và điều phối, còn ESP8266 cùng Arduino Uno thuộc lớp phần cứng.",
            "Luồng hoạt động bắt đầu từ camera DroidCam. Flask đọc từng khung hình, đưa vào YOLO để nhận diện, sau đó hiển thị kết quả lên web. Khi kết quả ổn định, server đưa lệnh vào hàng đợi. Worker nền lấy lệnh này và gửi HTTP đến ESP8266.",
            "ESP8266 nhận lệnh servo=1 hoặc servo=2, sau đó gửi ký tự tương ứng qua Serial cho Arduino Uno. Arduino Uno điều khiển servo ở chân D9 hoặc D10 để thực hiện phân loại."
        ]),
        ("2.3 Thiết kế dữ liệu và huấn luyện mô hình", [
            "Dữ liệu được quản lý trên Roboflow với hai lớp: Qua Hong và Qua Tuoi. Thứ tự lớp trong dataset rất quan trọng vì server ánh xạ class 0 thành quả hỏng và class 1 thành quả tươi.",
            "Sau khi tải dataset, file train_model.py dùng mô hình nền yolov8s.pt để huấn luyện lại. Kết quả huấn luyện được lưu tại runs/detect/rotten_fruit/weights/best.pt. Server chỉ sử dụng mô hình này, không fallback sang mô hình YOLO mặc định.",
            "Việc chỉ dùng mô hình đã huấn luyện riêng giúp chương trình tập trung vào bài toán của đề tài, tránh nhận diện các đối tượng không liên quan như người, đồ vật hoặc các lớp COCO mặc định."
        ]),
        ("2.4 Thiết kế server Flask", [
            "Server Flask được viết trong file server.py. Khi khởi động, server kiểm tra sự tồn tại của file best.pt, nạp mô hình YOLO và chuẩn bị các route web. Nếu không tìm thấy mô hình, chương trình báo lỗi để tránh chạy sai dữ liệu.",
            "Server có các chức năng chính: mở camera, xử lý nhận diện, ghi lịch sử, cung cấp luồng video, cung cấp trạng thái phần cứng và gửi lệnh điều khiển servo. Các thông số như IP DroidCam, IP ESP8266, ngưỡng confidence và thời gian cooldown có thể cấu hình bằng biến môi trường.",
            "Để tăng tính ổn định, server sử dụng queue cho lệnh servo. Queue chỉ giữ lệnh mới nhất, giúp tránh tồn đọng lệnh cũ khi phần cứng phản hồi chậm."
        ]),
        ("2.5 Thiết kế giao diện web", [
            "Giao diện web hiển thị video trực tiếp từ DroidCam, trong đó các quả được nhận diện sẽ có khung màu tương ứng. Quả tươi hiển thị màu xanh, quả hỏng hiển thị màu đỏ.",
            "Bên cạnh video, giao diện còn có phần thống kê số lượng quả tươi và quả hỏng, bảng lịch sử quét và khu vực trạng thái phần cứng. Khu vực phần cứng hiển thị URL ESP8266, lệnh gần nhất, phản hồi OK, lỗi kết nối, thời gian cooldown và số lệnh đang chờ.",
            "Nhờ phần trạng thái này, người dùng có thể biết lỗi đang nằm ở nhận diện, camera hay kết nối ESP8266 mà không cần xem terminal liên tục."
        ]),
        ("2.6 Thiết kế ESP8266", [
            "ESP8266 được lập trình để ưu tiên kết nối vào Wi-Fi chính. Nếu không kết nối được, ESP8266 bật Wi-Fi dự phòng có tên PHAN_LOAI_TRAI_CAY. Cách này giúp quá trình kiểm thử linh hoạt hơn trong trường hợp router hoặc mạng chính gặp vấn đề.",
            "ESP8266 cung cấp endpoint /ping để kiểm tra kết nối và endpoint /control để nhận lệnh điều khiển. Khi nhận /control?servo=1, ESP8266 gửi ký tự 1 qua Serial. Khi nhận /control?servo=2, ESP8266 gửi ký tự 2.",
            "Phản hồi HTTP từ ESP8266 được Flask ghi lại và hiển thị trên giao diện web. Nếu ESP8266 timeout, giao diện sẽ báo lỗi và server tạm dừng gửi lại trong một khoảng thời gian để chống spam."
        ]),
        ("2.7 Thiết kế Arduino Uno và servo", [
            "Arduino Uno nhận dữ liệu Serial ở tốc độ 9600 baud. Lệnh 1 điều khiển servoHong ở chân D9, còn lệnh 2 điều khiển servoTuoi ở chân D10. Sau khi servo quay đến góc gạt, Arduino đưa servo trở về góc ban đầu.",
            "Trong mô hình thực tế, servo quả tươi có thể cần một độ trễ nhất định để phù hợp với vị trí vật trên băng chuyền. Do đó code Arduino có hằng số SERVO_TUOI_DELAY để dễ hiệu chỉnh.",
            "Nguồn cấp cho servo nên là nguồn 5V ngoài, không nên lấy trực tiếp từ Arduino khi servo có tải lớn. Tất cả GND của Arduino, ESP8266 và nguồn servo phải nối chung."
        ]),
        ("2.8 Cơ chế chống spam và đảm bảo ổn định", [
            "Khi camera chạy liên tục, mô hình có thể nhận diện cùng một quả trong rất nhiều khung hình. Nếu mỗi khung hình đều gửi lệnh tới servo thì phần cứng sẽ bị spam. Vì vậy, hệ thống yêu cầu cùng một class xuất hiện ổn định trong một số frame liên tiếp mới phát lệnh.",
            "Mỗi servo có thời gian cooldown riêng. Nếu gửi lệnh thất bại, hệ thống cũng đặt thời gian chờ trước khi thử lại. Điều này giúp tránh việc log lỗi xuất hiện liên tục và tránh làm ảnh hưởng tới luồng DroidCam.",
            "Việc gửi HTTP tới ESP8266 được thực hiện trong worker nền. Nhờ đó, nếu ESP8266 phản hồi chậm hoặc mất kết nối, vòng đọc camera và YOLO vẫn tiếp tục chạy."
        ]),
    ]
    for title, paragraphs in chapter_2:
        add_section(document, title, paragraphs)
    heading(document, "2.9 Bảng ánh xạ lớp nhận diện và phần cứng", 2)
    table(document, ["Class", "Nhãn", "Trạng thái", "Lệnh ESP8266", "Chân Arduino"], [
        ["0", "Qua Hong", "Hỏng", "servo=1", "D9"],
        ["1", "Qua Tuoi", "Tươi", "servo=2", "D10"],
    ])
    document.add_page_break()

    heading(document, "CHƯƠNG 3. THỰC NGHIỆM, KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    chapter_3 = [
        ("3.1 Môi trường thực nghiệm", [
            "Hệ thống được triển khai trên máy tính Windows, sử dụng Python để chạy Flask server và mô hình YOLO. Phần cứng gồm ESP8266 NodeMCU, Arduino Uno và hai servo. Camera được lấy từ điện thoại thông qua DroidCam.",
            "Việc nạp code cho Arduino Uno và ESP8266 được thực hiện bằng PlatformIO trong VS Code. Project cung cấp các script hỗ trợ như upload_esp8266.ps1, upload_uno.ps1, monitor_esp8266.ps1 và run_all.ps1."
        ]),
        ("3.2 Kiểm thử mô hình nhận diện", [
            "Mô hình sau huấn luyện được kiểm tra bằng cách load file best.pt và xác nhận danh sách nhãn gồm Qua Hong và Qua Tuoi. Khi chạy thử với ảnh mẫu, mô hình trả về bounding box và confidence tương ứng.",
            "Kết quả ban đầu cho thấy mô hình có thể nhận diện hai lớp của đề tài. Tuy nhiên, do số lượng ảnh trong dataset chưa lớn, cần bổ sung thêm dữ liệu trong nhiều điều kiện ánh sáng, góc chụp và nền khác nhau để tăng độ ổn định."
        ]),
        ("3.3 Kiểm thử DroidCam", [
            "DroidCam được kiểm tra bằng cách mở luồng video từ IP đã cấu hình. Nếu không đọc được frame, server sẽ hiển thị frame thông báo lỗi thay vì làm web dừng đột ngột.",
            "Trong quá trình tích hợp phần cứng, việc gửi lệnh ESP8266 từng làm DroidCam bị đứng khi request timeout. Lỗi này được khắc phục bằng cách tách gửi lệnh sang worker nền, giúp luồng camera tiếp tục hoạt động."
        ]),
        ("3.4 Kiểm thử ESP8266", [
            "ESP8266 được kiểm tra bằng endpoint /ping. Nếu truy cập /ping nhận được phản hồi ESP8266 OK, có thể xác định ESP8266 đã sẵn sàng nhận lệnh. Sau đó kiểm tra /control?servo=1 và /control?servo=2 để xác nhận lệnh điều khiển.",
            "Các lỗi thường gặp gồm sai IP, khác mạng Wi-Fi, chưa nạp được firmware, cổng COM bị giữ bởi Serial Monitor hoặc cáp USB không truyền dữ liệu. Project đã bổ sung script kiểm tra và hướng dẫn xử lý các lỗi này."
        ]),
        ("3.5 Kiểm thử Arduino Uno và servo", [
            "Arduino Uno được kiểm thử riêng bằng cách gửi ký tự 1 hoặc 2 qua Serial Monitor. Nếu servo ở D9 và D10 quay đúng, phần Arduino hoạt động. Sau đó mới kết nối ESP8266 để kiểm tra luồng ESP8266 gửi lệnh cho Arduino.",
            "Khi servo không quay, cần kiểm tra nguồn cấp, dây tín hiệu, GND chung và việc nạp đúng chương trình cho Arduino. Nếu dùng nguồn yếu, servo có thể rung hoặc không hoạt động."
        ]),
        ("3.6 Kiểm thử toàn hệ thống", [
            "Sau khi từng module hoạt động, hệ thống được chạy bằng run_all.ps1. Người dùng mở giao diện web, đưa trái cây vào vùng camera và quan sát kết quả. Khi phát hiện Qua Hong, server gửi servo=1. Khi phát hiện Qua Tuoi, server gửi servo=2.",
            "Giao diện phần cứng cho biết lệnh gần nhất, phản hồi từ ESP8266 và lỗi nếu có. Đây là công cụ quan trọng để kiểm tra chương trình trong quá trình demo."
        ]),
        ("3.7 Kết quả đạt được", [
            "Project đã hoàn thành hệ thống nhận diện quả tươi và quả hỏng bằng YOLO, hiển thị kết quả trên web Flask, sử dụng DroidCam làm nguồn camera và tích hợp ESP8266 với Arduino Uno để điều khiển servo.",
            "Hệ thống có cơ chế chống spam lệnh, worker nền cho phần cứng, script hỗ trợ nạp code và giao diện hiển thị trạng thái phản hồi. Đây là nền tảng phù hợp cho một mô hình AIoT phục vụ phân loại nông sản."
        ]),
        ("3.8 Hạn chế và hướng phát triển", [
            "Hạn chế lớn nhất của hệ thống là dataset còn ít ảnh và chưa có tập kiểm thử độc lập đầy đủ. Vì vậy, độ chính xác thực tế có thể thay đổi khi gặp điều kiện ánh sáng hoặc loại quả khác với dữ liệu huấn luyện.",
            "Trong tương lai, có thể mở rộng dữ liệu, bổ sung camera ổn định hơn, cải tiến cơ cấu băng chuyền, thêm cảm biến phát hiện vật và lưu dữ liệu thống kê vào cơ sở dữ liệu để theo dõi lâu dài."
        ]),
    ]
    for title, paragraphs in chapter_3:
        add_section(document, title, paragraphs)
    heading(document, "3.9 Bảng tổng hợp lỗi và cách xử lý", 2)
    table(document, ["Lỗi", "Nguyên nhân", "Cách xử lý"], [
        ["ESP8266 timeout", "Sai IP hoặc khác mạng", "Kiểm tra /ping, dùng auto-discover hoặc nhập đúng IP"],
        ["DroidCam bị đứng", "Request phần cứng chặn luồng video", "Dùng queue và worker nền"],
        ["Upload ESP8266 thất bại", "Board chưa vào bootloader hoặc COM bị giữ", "Đóng Serial Monitor, tháo RX/TX, giữ FLASH/BOOT nếu cần"],
        ["Servo không quay", "Nguồn yếu hoặc sai dây", "Dùng nguồn 5V ngoài, nối chung GND và kiểm tra chân D9/D10"],
    ])
    document.add_page_break()

    heading(document, "KẾT LUẬN", 1)
    para(document, "Đề tài đã xây dựng được một hệ thống mẫu kết hợp giữa trí tuệ nhân tạo và điều khiển phần cứng. Hệ thống có thể nhận diện quả tươi và quả hỏng bằng YOLO, hiển thị kết quả trên web, đồng thời gửi lệnh tới ESP8266 và Arduino Uno để điều khiển servo phân loại.")
    para(document, "Mặc dù còn hạn chế về dữ liệu và cơ khí, project đã thể hiện đầy đủ quy trình xây dựng một ứng dụng AIoT từ dữ liệu, mô hình, giao diện, giao tiếp mạng đến hành động vật lý của servo.")
    document.add_page_break()

    heading(document, "TÀI LIỆU THAM KHẢO", 1)
    for item in [
        "Tài liệu Ultralytics YOLO.",
        "Tài liệu Flask Framework.",
        "Tài liệu OpenCV.",
        "Tài liệu Roboflow.",
        "Tài liệu ESP8266 Arduino Core.",
        "Tài liệu thư viện Servo của Arduino.",
    ]:
        bullet(document, item)
    document.add_page_break()

    heading(document, "PHỤ LỤC", 1)
    code_excerpt(document, "Phụ lục A. Trích đoạn server.py", "server.py")
    code_excerpt(document, "Phụ lục B. Trích đoạn ESP8266", "src/esp8266_node/main.cpp")
    code_excerpt(document, "Phụ lục C. Trích đoạn Arduino Uno", "src/arduino_uno/main.cpp")
    code_excerpt(document, "Phụ lục D. Trích đoạn run_all.ps1", "run_all.ps1")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
