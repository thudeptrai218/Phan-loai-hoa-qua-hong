from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def clear_document(document):
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_normal_style(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)


def add_title(document, text, size=18):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
    return heading


def add_paragraph(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Inches(0.3)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_bullets(document, items):
    for item in items:
        try:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(item)
        except KeyError:
            add_paragraph(document, "- " + item)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def add_code_block(document, lines):
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_page(document, title, paragraphs, bullets=None, table=None):
    add_heading(document, title, level=1)
    for paragraph in paragraphs:
        add_paragraph(document, paragraph)
    if bullets:
        add_bullets(document, bullets)
    if table:
        add_table(document, table["headers"], table["rows"])
    document.add_page_break()


def read_excerpt(path, max_lines=28):
    file_path = Path(path)
    if not file_path.exists():
        return [f"Khong tim thay file: {path}"]
    lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
    return lines[:max_lines]


def main():
    if len(sys.argv) < 3:
        raise SystemExit("Usage: python generate_project_report.py TEMPLATE_DOCX OUTPUT_DOCX")

    template_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not template_path.exists():
        raise FileNotFoundError(template_path)

    document = Document(str(template_path))
    clear_document(document)
    set_normal_style(document)

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)

    add_title(document, "BAO CAO DO AN", 18)
    add_title(document, "HE THONG PHAN LOAI QUA TUOI VA QUA HONG", 18)
    add_title(document, "SU DUNG YOLO, DROIDCAM, FLASK, ESP8266 VA ARDUINO UNO", 14)
    document.add_paragraph()
    add_title(document, "Sinh vien thuc hien: Doan Tuan Nam", 13)
    add_title(document, "Thoi gian tao bao cao: " + datetime.now().strftime("%d/%m/%Y %H:%M"), 12)
    document.add_page_break()

    add_heading(document, "MUC LUC", 1)
    add_paragraph(document, "1. Tong quan de tai")
    add_paragraph(document, "2. Co so ly thuyet va cong nghe su dung")
    add_paragraph(document, "3. Phan tich yeu cau va thiet ke he thong")
    add_paragraph(document, "4. Xay dung tap du lieu va train model")
    add_paragraph(document, "5. Xay dung ung dung web va nhan dien thoi gian thuc")
    add_paragraph(document, "6. Tich hop ESP8266, Arduino Uno va servo")
    add_paragraph(document, "7. Kiem thu, danh gia va huong phat trien")
    add_paragraph(document, "8. Phu luc ma nguon")
    document.add_page_break()

    pages = [
        (
            "1.1 Ly do chon de tai",
            [
                "Phan loai trai cay theo tinh trang tuoi va hong la mot bai toan co tinh ung dung cao trong nong nghiep, bao quan thuc pham va cac day chuyen phan loai tu dong. Neu thuc hien bang thu cong, viec kiem tra phu thuoc vao kinh nghiem cua nguoi lao dong, toc do khong on dinh va de xay ra sai sot khi so luong san pham lon.",
                "De tai nay ket hop thi giac may tinh voi phan cung dieu khien. Camera cung cap hinh anh trai cay, mo hinh YOLO nhan dien qua tuoi va qua hong, ung dung Flask hien thi ket qua tren web, con ESP8266 va Arduino Uno thuc hien dieu khien servo de gat san pham ve dung khay.",
                "Viec thuc hien de tai giup sinh vien nam duoc mot quy trinh hoan chinh tu thu thap du lieu, gan nhan, huan luyen mo hinh, xay dung giao dien web cho den tich hop phan cung dieu khien."
            ],
        ),
        (
            "1.2 Muc tieu cua de tai",
            [
                "Muc tieu chinh cua de tai la xay dung mot he thong co kha nang nhan dien trai cay theo hai trang thai: qua tuoi va qua hong. Ket qua nhan dien duoc hien thi truc tiep tren web va duoc su dung de dieu khien servo thong qua ESP8266 va Arduino Uno.",
                "He thong can dam bao cac yeu cau co ban: camera cap du lieu lien tuc, mo hinh nhan dien dung hai nhan, web hien thi lich su va trang thai phan cung, ESP8266 nhan lenh HTTP, Arduino Uno nhan lenh Serial va dieu khien hai servo."
            ],
            {
                "bullets": [
                "Nhan dien duoc class Qua Hong va Qua Tuoi.",
                "Hien thi video truc tiep tren giao dien web.",
                "Gui lenh servo=1 cho qua hong va servo=2 cho qua tuoi.",
                "Co co che chong spam lenh khi nhan dien lien tuc.",
                "Co trang thai phan hoi ESP8266 tren giao dien."
                ],
            },
        ),
        (
            "1.3 Pham vi thuc hien",
            [
                "Pham vi de tai tap trung vao mo hinh thu nghiem quy mo nho. Camera duoc lay tu DroidCam, model duoc train bang du lieu Roboflow, server chay tren laptop, ESP8266 dam nhan vai tro cau noi Wi-Fi va Arduino Uno dieu khien servo.",
                "He thong chua tap trung vao thiet ke co khi cong nghiep hoan chinh, ma uu tien minh hoa duoc day du luong xu ly du lieu tu camera den hanh dong vat ly cua servo."
            ],
        ),
        (
            "1.4 Ket qua mong doi",
            [
                "Sau khi hoan thanh, nguoi dung co the chay server Flask, mo giao dien web, dua trai cay vao khung hinh camera va quan sat ket qua nhan dien. Neu phat hien qua hong, he thong gui lenh toi ESP8266 va Arduino Uno de servo thu nhat gat qua. Neu phat hien qua tuoi, servo thu hai se duoc kich hoat.",
                "Ngoai ra, giao dien web can cho biet tinh trang giao tiep voi phan cung, bao gom lenh gan nhat, phan hoi OK tu ESP8266 va loi neu ket noi khong thanh cong."
            ],
        ),
        (
            "2.1 Tong quan ve tri tue nhan tao trong nhan dien anh",
            [
                "Tri tue nhan tao trong thi giac may tinh cho phep may tinh phan tich hinh anh va rut ra thong tin co y nghia. Trong de tai nay, AI duoc ung dung de phat hien doi tuong trai cay va phan loai tinh trang cua trai cay theo nhan Qua Hong hoac Qua Tuoi.",
                "Khac voi cac phuong phap xu ly anh truyen thong dua tren nguong mau hoac dac trung thu cong, mo hinh hoc sau co kha nang hoc dac trung tu du lieu. Dieu nay phu hop voi bai toan qua hong vi dau hieu hong co the khac nhau tuy loai qua, anh sang va goc chup."
            ],
        ),
        (
            "2.2 Mo hinh YOLO",
            [
                "YOLO la mo hinh phat hien doi tuong thoi gian thuc. Ten YOLO la viet tat cua You Only Look Once, the hien cach mo hinh du doan vi tri va lop doi tuong trong mot lan xu ly anh.",
                "Trong project, YOLO duoc su dung thong qua thu vien Ultralytics. Model co dau vao la frame camera va dau ra la cac hop gioi han, do tin cay va class id. Class id duoc anh xa thanh Qua Hong hoac Qua Tuoi de hien thi va dieu khien servo."
            ],
        ),
        (
            "2.3 Roboflow va xay dung dataset",
            [
                "Roboflow duoc su dung de quan ly du lieu anh, gan nhan va export dataset theo dinh dang YOLOv8. Dataset moi cua project co hai class: Qua Hong va Qua Tuoi. Qua trinh export tao ra cau truc thu muc train/images, train/labels va file data.yaml.",
                "Do dataset hien tai con nho, viec mo rong them anh trong nhieu dieu kien anh sang, goc chup va nen khac nhau la can thiet de model hoat dong on dinh hon trong thuc te."
            ],
        ),
        (
            "2.4 Flask trong ung dung web",
            [
                "Flask la framework Python nhe dung de xay dung web server. Trong project, Flask dam nhan viec phuc vu giao dien web, tao luong video MJPEG, tra ve lich su nhan dien va trang thai phan cung.",
                "Cac route chinh gom trang chu, video_feed, history va hardware_status. Route video_feed lay frame da ve bounding box tu YOLO va gui lien tuc ve trinh duyet."
            ],
        ),
        (
            "2.5 DroidCam lam nguon camera",
            [
                "DroidCam bien dien thoai thanh camera IP cho may tinh. Project cau hinh mac dinh dia chi camera la 172.20.10.3 va port 4747. Server thu cac endpoint thong dung nhu /video, /mjpegfeed va /videofeed de tang kha nang ket noi.",
                "Su dung DroidCam giup tan dung camera dien thoai co chat luong tot hon webcam co san, dong thoi linh hoat trong viec dat camera len mo hinh bang chuyen."
            ],
        ),
        (
            "2.6 ESP8266",
            [
                "ESP8266 la vi dieu khien co Wi-Fi tich hop. Trong project, ESP8266 dong vai tro cau noi giua Flask server va Arduino Uno. Flask gui yeu cau HTTP den ESP8266, sau do ESP8266 gui ky tu 1 hoac 2 qua Serial cho Arduino.",
                "Firmware ESP8266 duoc cau hinh de co gang ket noi Wi-Fi chinh. Neu khong thanh cong, ESP8266 co the bat Wi-Fi du phong PHAN_LOAI_TRAI_CAY de phuc vu qua trinh test."
            ],
        ),
        (
            "2.7 Arduino Uno va servo",
            [
                "Arduino Uno nhan lenh tu ESP8266 qua cong Serial 9600. Khi nhan lenh 1, Arduino kich hoat servo gat qua hong o chan D9. Khi nhan lenh 2, Arduino kich hoat servo gat qua tuoi o chan D10.",
                "Servo duoc dieu khien bang thu vien Servo.h. Vi servo co the can dong lon, nguon servo nen la nguon 5V ngoai va phai noi chung GND voi Arduino."
            ],
        ),
        (
            "3.1 Yeu cau chuc nang",
            [
                "He thong can co kha nang nhan video tu DroidCam, xu ly tung frame bang YOLO, ve hop nhan dien len anh, hien thi ket qua tren web va luu lich su nhan dien. Khi co ket qua on dinh, he thong gui lenh dieu khien servo toi ESP8266.",
                "Chuong trinh cung can co che thong bao loi neu camera khong ket noi duoc hoac ESP8266 khong phan hoi. Nhung loi nay duoc hien thi tren giao dien de nguoi dung de dang debug."
            ],
            {
                "table": {
                    "headers": ["Thanh phan", "Chuc nang"],
                    "rows": [
                        ["DroidCam", "Cung cap video dau vao"],
                        ["YOLO", "Nhan dien Qua Hong va Qua Tuoi"],
                        ["Flask", "Hien thi web va dieu phoi xu ly"],
                        ["ESP8266", "Nhan HTTP va gui Serial"],
                        ["Arduino Uno", "Dieu khien hai servo"],
                    ],
                },
            },
        ),
        (
            "3.2 Yeu cau phi chuc nang",
            [
                "He thong can hoat dong on dinh trong thoi gian dai, khong bi treo khi ESP8266 mat ket noi, khong gui lenh servo lien tuc gay qua tai phan cung va co kha nang cau hinh cac dia chi IP bang bien moi truong hoac script PowerShell.",
                "De tang do on dinh, project da tach viec gui lenh ESP8266 ra worker nen. Luong camera va YOLO khong bi chan boi request HTTP toi phan cung."
            ],
        ),
        (
            "3.3 Kien truc tong the",
            [
                "Kien truc cua he thong gom bon lop: lop thu nhan du lieu, lop xu ly AI, lop giao dien va lop dieu khien phan cung. DroidCam thuoc lop thu nhan du lieu, YOLO thuoc lop xu ly AI, Flask web thuoc lop giao dien, con ESP8266 va Arduino Uno thuoc lop dieu khien phan cung.",
                "Cach chia lop nay giup de debug tung phan. Co the chay web voi che do SkipHardware de test model va camera truoc, sau do moi bat phan cung."
            ],
        ),
        (
            "3.4 Luong xu ly khi phat hien trai cay",
            [
                "Moi frame camera duoc dua vao model YOLO. Neu ket qua co confidence dat nguong va class thuoc tap hop hop le, server ve hop len frame va them vao lich su. Neu cung mot class xuat hien on dinh trong so frame cau hinh, server day lenh vao hang doi servo.",
                "Worker nen lay lenh trong hang doi va gui HTTP den ESP8266. Neu ESP8266 phan hoi OK, giao dien cap nhat trang thai thanh cong. Neu timeout, he thong luu loi va tam dung gui lai trong mot khoang thoi gian."
            ],
        ),
        (
            "3.5 Co che chong spam lenh servo",
            [
                "Nhan dien thoi gian thuc co the tao ra nhieu ket qua moi giay. Neu moi ket qua deu gui lenh servo, phan cung se bi spam, servo rung lien tuc va ESP8266 co the qua tai. Do do project ap dung nhieu lop chong spam.",
                "Thu nhat, class phai on dinh trong mot so frame lien tiep. Thu hai, moi servo co cooldown rieng. Thu ba, neu gui that bai, he thong tam dung gui trong vai giay. Thu tu, queue chi giu lenh moi nhat de tranh ton lenh cu."
            ],
        ),
        (
            "4.1 Cau truc dataset Roboflow",
            [
                "Dataset Roboflow duoc tai ve trong thu muc traicay-1. File data_train_only.yaml duoc tao de train tam bang thu muc train khi dataset chua co valid/test. Dataset co hai class theo thu tu Roboflow: 0 la Qua Hong va 1 la Qua Tuoi.",
                "Thu tu class nay rat quan trong vi server anh xa class 0 thanh servo qua hong va class 1 thanh servo qua tuoi."
            ],
            {
                "table": {
                    "headers": ["Class ID", "Nhan", "Hanh dong"],
                    "rows": [["0", "Qua Hong", "Gui servo=1"], ["1", "Qua Tuoi", "Gui servo=2"]],
                },
            },
        ),
        (
            "4.2 Qua trinh tai dataset",
            [
                "Project co script download_roboflow_dataset.py de tai dataset tu Roboflow bang API key. Nguoi dung cau hinh ROBOFLOW_API_KEY va ROBOFLOW_VERSION, sau do chay script de tai du lieu ve may.",
                "Script nay giup giam thao tac thu cong va dam bao dataset duoc dat trong project voi cau truc phu hop cho train_model.py."
            ],
        ),
        (
            "4.3 Qua trinh train model",
            [
                "File train_model.py su dung YOLO base model yolov8s.pt va train tren data.yaml da chon. Ket qua duoc luu vao runs/detect/rotten_fruit/weights/best.pt. Server chi load file best.pt nay va khong fallback sang model mac dinh de tranh nhan dien nham doi tuong khac.",
                "Qua trinh train da hoan thanh tren dataset Roboflow version 1. Model load duoc hai nhan Qua Hong va Qua Tuoi."
            ],
        ),
        (
            "4.4 Danh gia dataset hien tai",
            [
                "Dataset hien tai co so luong anh con it, vi vay do chinh xac tren thuc te co the chua cao. Cac chi so validation khi dung chung train lam val chi mang tinh tham khao vi chua co tap kiem thu doc lap.",
                "De nang cao chat luong, can bo sung anh qua tuoi va qua hong trong nhieu dieu kien: anh sang manh, anh sang yeu, nen toi, nen sang, goc gan, goc xa, nhieu loai trai cay va cac muc do hong khac nhau."
            ],
        ),
        (
            "5.1 Server Flask",
            [
                "Server Flask la trung tam dieu phoi cua he thong. Khi chay server.py, Flask khoi tao model YOLO, mo nguon camera, phuc vu giao dien web va tao cac route API. Route chinh tra ve template web_dashboard.html, route video_feed tra ve luong video, route history tra ve lich su va route hardware_status tra ve trang thai ESP8266.",
                "Cac thong so nhu DroidCam IP, ESP8266 IP, nguong confidence va cooldown servo co the cau hinh bang bien moi truong thong qua run_all.ps1."
            ],
        ),
        (
            "5.2 Xu ly video DroidCam",
            [
                "DroidCam duoc doc bang OpenCV VideoCapture. He thong thu nhieu endpoint de tang kha nang ket noi. Neu khong doc duoc frame, server tao frame thong bao loi thay vi lam web dung dot ngot.",
                "De tranh treo camera khi ESP8266 cham, viec gui lenh phan cung da duoc tach ra worker nen."
            ],
        ),
        (
            "5.3 Giao dien web",
            [
                "Giao dien web hien thi video truc tiep, legend mau xanh/do, thong ke so qua tuoi va qua hong, bang lich su quet va trang thai phan cung. Trang thai phan cung gom URL ESP8266, lenh gan nhat, phan hoi OK, loi va so lenh dang cho.",
                "Giao dien duoc thiet ke de phuc vu qua trinh demo va debug. Khi ESP8266 mat ket noi, nguoi dung co the nhin thay loi ngay tren web."
            ],
        ),
        (
            "5.4 API lich su va trang thai phan cung",
            [
                "API /history tra ve danh sach cac lan nhan dien gan nhat. Moi ban ghi gom thoi gian, ten doi tuong va trang thai. API /hardware_status tra ve thong tin ve phan cung, giup giao dien cap nhat dinh ky ma khong can reload trang.",
                "Cach thiet ke nay tach rieng giao dien va logic server, giup sau nay co the thay doi giao dien ma khong anh huong den xu ly backend."
            ],
        ),
        (
            "6.1 Thiet ke giao tiep Flask va ESP8266",
            [
                "Flask gui lenh den ESP8266 thong qua HTTP GET. Dia chi co dang /control?servo=1 hoac /control?servo=2. ESP8266 kiem tra tham so servo, neu hop le thi in ky tu 1 hoac 2 ra Serial va tra ve phan hoi OK.",
                "Cach giao tiep HTTP don gian, de test bang trinh duyet va PowerShell. Endpoint /ping duoc bo sung de Flask hoac nguoi dung kiem tra ESP8266 da san sang hay chua."
            ],
        ),
        (
            "6.2 Thiet ke giao tiep ESP8266 va Arduino Uno",
            [
                "ESP8266 va Arduino Uno giao tiep qua Serial 9600. ESP8266 TX noi voi Arduino RX D0, Arduino TX noi voi ESP8266 RX qua mach chia ap. Tuy nhien trong project nay Arduino chu yeu can nhan lenh tu ESP8266, nen day ESP8266 TX sang Arduino RX la quan trong nhat.",
                "Can noi chung GND giua ESP8266 va Arduino. Khi upload code cho Arduino, nen thao day D0/D1 de tranh loi nap."
            ],
        ),
        (
            "6.3 Dieu khien servo bang Arduino",
            [
                "Arduino dung thu vien Servo.h de dieu khien hai servo. Servo qua hong gan chan D9, servo qua tuoi gan chan D10. Khi nhan lenh 1, servo qua hong quay den goc gat va sau do ve lai goc ban dau. Khi nhan lenh 2, servo qua tuoi co do tre de phu hop vi tri bang chuyen.",
                "Thong so goc quay va thoi gian tre co the dieu chinh trong file servo.ino hoac src/arduino_uno/main.cpp."
            ],
        ),
        (
            "6.4 Dau day phan cung",
            [
                "Servo can nguon 5V ngoai de dam bao dong cap du. Day tin hieu cua servo qua hong noi vao D9, servo qua tuoi noi vao D10. GND cua nguon servo phai noi voi GND Arduino de tin hieu dieu khien co cung moc dien ap.",
                "ESP8266 hoat dong o muc logic 3.3V. Arduino Uno TX la 5V nen neu noi nguoc ve ESP8266 RX can mach chia ap. Neu chi can ESP8266 gui lenh sang Arduino, co the uu tien noi ESP8266 TX sang Arduino RX va GND chung."
            ],
        ),
        (
            "6.5 Che do ket noi ESP8266",
            [
                "Firmware ESP8266 duoc cau hinh de uu tien ket noi Wi-Fi chinh. Neu that bai, ESP8266 bat Wi-Fi du phong PHAN_LOAI_TRAI_CAY. Cach nay giup test doc lap khi router khong cap IP hoac laptop khong scan duoc ESP8266.",
                "Trong van hanh on dinh voi DroidCam qua Wi-Fi, nen de ESP8266, laptop va dien thoai cung mot mang Wi-Fi. Neu laptop ket noi vao Wi-Fi du phong cua ESP8266, DroidCam qua router co the bi mat ket noi."
            ],
        ),
        (
            "7.1 Kiem thu tung thanh phan",
            [
                "Kiem thu nen thuc hien theo thu tu: kiem tra model nhan dien bang anh mau, kiem tra DroidCam co tra video, kiem tra server Flask hien thi web, kiem tra ESP8266 bang /ping, kiem tra servo bang /control?servo=1 va /control?servo=2, cuoi cung moi ghep toan bo he thong.",
                "Cach kiem thu tung thanh phan giup xac dinh dung vi tri loi thay vi phan doan toan he thong."
            ],
        ),
        (
            "7.2 Cac loi da gap va cach khac phuc",
            [
                "Trong qua trinh phat trien, mot so loi da xay ra nhu ESP8266 khong truy cap duoc do sai subnet, PlatformIO khong co PATH, ESP8266 khong vao bootloader, cong COM bi semaphore timeout va DroidCam bi dung khi request ESP8266 timeout.",
                "Cac loi nay da duoc khac phuc bang cach cai PlatformIO CLI, cau hinh PLATFORMIO_CORE_DIR trong project, giam upload speed, them script upload rieng, them worker nen cho lenh ESP8266 va them trang thai phan cung tren web."
            ],
        ),
        (
            "7.3 Danh gia uu diem",
            [
                "He thong co uu diem la kien truc ro rang, moi thanh phan co the test rieng, giao dien web co trang thai phan cung, model khong fallback sang nhan khong lien quan, va co script mot lenh de chay chuong trinh.",
                "Viec su dung Roboflow giup qua trinh cap nhat dataset va train lai model thuan tien hon so voi quan ly thu cong."
            ],
        ),
        (
            "7.4 Han che",
            [
                "Han che lon nhat la dataset con nho nen kha nang tong quat cua model chua cao. He thong cung phu thuoc vao do on dinh cua mang Wi-Fi neu DroidCam va ESP8266 deu dung ket noi khong day.",
                "Ngoai ra, mo hinh co khi bang chuyen va servo moi o muc demo, can tinh toan them vi tri dat camera, khoang cach servo va thoi gian tre phu hop toc do bang chuyen."
            ],
        ),
        (
            "7.5 Huong phat trien",
            [
                "Trong tuong lai co the mo rong dataset, bo sung valid/test doc lap, cai thien thiet ke co khi, them cam bien phat hien vat tren bang chuyen, thay DroidCam bang camera cong nghiep hoac camera USB on dinh hon.",
                "Co the bo sung che do cau hinh Wi-Fi ESP8266 qua web, luu lich su vao co so du lieu va hien thi thong ke theo ngay."
            ],
        ),
    ]

    for title, paragraphs, *rest in pages:
        options = rest[0] if rest else {}
        if isinstance(options, dict):
            add_page(document, title, paragraphs, options.get("bullets"), options.get("table"))
        else:
            add_page(document, title, paragraphs)

    code_sections = [
        ("8.1 Phu luc server.py", "server.py"),
        ("8.2 Phu luc train_model.py", "train_model.py"),
        ("8.3 Phu luc ESP8266", "src/esp8266_node/main.cpp"),
        ("8.4 Phu luc Arduino Uno", "src/arduino_uno/main.cpp"),
        ("8.5 Phu luc run_all.ps1", "run_all.ps1"),
    ]

    for title, path in code_sections:
        add_heading(document, title, level=1)
        add_paragraph(document, f"Doan ma duoi day la trich doan tu file {path}, duoc dua vao bao cao de minh hoa cau truc cai dat cua project.")
        add_code_block(document, read_excerpt(path))
        document.add_page_break()

    add_heading(document, "KET LUAN", level=1)
    add_paragraph(document, "De tai da xay dung duoc mot he thong mau ket hop giua nhan dien anh va dieu khien phan cung. He thong co the nhan dien qua tuoi va qua hong bang YOLO, hien thi ket qua tren web Flask, dong thoi gui lenh toi ESP8266 va Arduino Uno de dieu khien servo phan loai.")
    add_paragraph(document, "Mac du con han che ve quy mo dataset va co khi, project da the hien day du quy trinh xay dung mot ung dung AIoT: tu du lieu, mo hinh, giao dien, giao tiep mang den dieu khien thiet bi vat ly.")
    document.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
